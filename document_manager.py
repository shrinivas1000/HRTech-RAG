from fastapi import UploadFile, HTTPException
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
import tempfile
import os
import datetime
import json
import shutil
import pickle
from pathlib import Path
from typing import List, Dict
import logging
from docx import Document as DocxDocument

class DocumentManager:
    def __init__(self, vectorstore_path: str = "vectorstore", data_path: str = "data"):
        self.vectorstore_path = Path(vectorstore_path)
        self.data_path = Path(data_path)
        self.data_path.mkdir(exist_ok=True)
        self.vectorstore_path.mkdir(exist_ok=True)
        
        # FAISS-specific file paths
        self.faiss_index_path = self.vectorstore_path / "index.faiss"
        self.faiss_metadata_path = self.vectorstore_path / "index.pkl"
        
        self.tracking_file = self.data_path / "document_tracking.json"
        self.uploaded_docs = self._load_document_tracking()
        
    def _load_document_tracking(self) -> List[Dict]:
        """Load document tracking from persistent storage"""
        try:
            if self.tracking_file.exists():
                with open(self.tracking_file, 'r') as f:
                    return json.load(f)
        except Exception as e:
            logging.warning(f"Failed to load document tracking: {e}")
        return []
    
    def _save_document_tracking(self) -> None:
        """Save document tracking to persistent storage"""
        try:
            with open(self.tracking_file, 'w') as f:
                json.dump(self.uploaded_docs, f, indent=2)
        except Exception as e:
            logging.error(f"Failed to save document tracking: {e}")
    
    def _validate_file_type(self, filename: str) -> str:
        """Validate and return file type"""
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return 'pdf'
        elif filename_lower.endswith('.docx'):
            return 'docx'
        else:
            raise HTTPException(
                status_code=400, 
                detail="Only PDF and DOCX files allowed"
            )
    
    async def store_uploaded_file(self, file: UploadFile) -> Dict:
        """Store file and automatically rebuild embeddings"""
        
        # Validate file type
        file_type = self._validate_file_type(file.filename)
        
        # Generate unique filename to prevent conflicts
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        file_extension = '.pdf' if file_type == 'pdf' else '.docx'
        base_name = Path(file.filename).stem
        safe_filename = f"{timestamp}_{base_name}{file_extension}"
        file_path = self.data_path / safe_filename
        
        # Check for duplicates by original filename
        for doc in self.uploaded_docs:
            if doc["filename"] == file.filename:
                raise HTTPException(
                    status_code=409, 
                    detail=f"Document with name '{file.filename}' already exists"
                )
        
        # Save file to data/ folder
        try:
            content = await file.read()
            
            # Validate file size (50MB limit)
            if len(content) > 50 * 1024 * 1024:
                raise HTTPException(status_code=413, detail="File too large (max 50MB)")
            
            with open(file_path, "wb") as f:
                f.write(content)
            
            # Track document metadata
            doc_info = {
                "filename": file.filename,
                "stored_filename": safe_filename,
                "file_path": str(file_path),
                "file_type": file_type,
                "upload_date": datetime.datetime.now().isoformat(),
                "file_size": len(content),
                "status": "stored",
                "chunk_count": 0,
                "last_processed": None
            }
            
            self.uploaded_docs.append(doc_info)
            self._save_document_tracking()
            
            # AUTO-PROCESS: Rebuild embeddings immediately
            try:
                processing_result = self.process_all_documents()
                logging.info(f"Auto-processed after upload: {processing_result}")
                doc_info["auto_processed"] = True
            except Exception as e:
                logging.error(f"Auto-processing failed: {e}")
                doc_info["auto_processed"] = False
            
            logging.info(f"Stored {file_type.upper()} file: {file.filename} as {safe_filename}")
            return doc_info
            
        except HTTPException:
            raise
        except Exception as e:
            # Clean up file if it was created
            if file_path.exists():
                file_path.unlink()
            logging.error(f"Failed to store file {file.filename}: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to store file: {str(e)}")
    
    def remove_document(self, stored_filename: str) -> bool:
        """Remove document and automatically rebuild embeddings"""
        try:
            # Find document in tracking
            doc_to_remove = None
            for doc in self.uploaded_docs:
                if doc["stored_filename"] == stored_filename:
                    doc_to_remove = doc
                    break
            
            if not doc_to_remove:
                return False
            
            # Remove physical file
            file_path = Path(doc_to_remove["file_path"])
            if file_path.exists():
                file_path.unlink()
                logging.info(f"Removed file: {file_path}")
            
            # Remove from tracking
            self.uploaded_docs.remove(doc_to_remove)
            self._save_document_tracking()
            
            # AUTO-PROCESS: Rebuild embeddings immediately
            try:
                processing_result = self.process_all_documents()
                logging.info(f"Auto-processed after removal: {processing_result}")
            except Exception as e:
                logging.error(f"Auto-processing failed: {e}")
            
            return True
            
        except Exception as e:
            logging.error(f"Failed to remove document {stored_filename}: {e}")
            return False
    
    def _clear_vectorstore(self):
        """Clear FAISS vector store - much simpler than ChromaDB"""
        print(f"🔍 DEBUG: Clearing FAISS vector store")
        
        try:
            # Simply remove the FAISS files
            if self.faiss_index_path.exists():
                self.faiss_index_path.unlink()
                print(f"🗑️ DEBUG: Removed FAISS index file")
            
            if self.faiss_metadata_path.exists():
                self.faiss_metadata_path.unlink()
                print(f"🗑️ DEBUG: Removed FAISS metadata file")
            
            # Remove any other files in vectorstore directory
            for file in self.vectorstore_path.iterdir():
                if file.is_file():
                    file.unlink()
                    print(f"🗑️ DEBUG: Removed additional file: {file.name}")
            
            print(f"✅ DEBUG: FAISS vector store cleared successfully")
            
        except Exception as e:
            print(f"❌ DEBUG: Failed to clear FAISS vector store: {e}")
            logging.error(f"Failed to clear FAISS vector store: {e}")
            raise
    
    def _process_pdf(self, file_path: str, filename: str) -> List[Document]:
        """Process PDF file into chunks"""
        loader = PyPDFLoader(file_path)
        pages = loader.load_and_split()
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        documents = []
        for page in pages:
            chunks = text_splitter.split_text(page.page_content)
            for chunk_idx, chunk in enumerate(chunks):
                chunk_metadata = {
                    "page": page.metadata.get("page", 0),
                    "source": filename,
                    "file_type": "pdf",
                    "chunk_index": chunk_idx,
                    "upload_date": datetime.datetime.now().isoformat()
                }
                
                documents.append(Document(
                    page_content=chunk,
                    metadata=chunk_metadata
                ))
        
        return documents
    
    def _process_docx(self, file_path: str, filename: str) -> List[Document]:
        """Process DOCX file into chunks"""
        try:
            doc = DocxDocument(file_path)
            
            # Extract text from all paragraphs
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Join paragraphs with double newlines
            content = "\n\n".join(full_text)
            
            if not content.strip():
                raise ValueError("No readable content found in DOCX file")
            
            # Split into chunks
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                separators=["\n\n", "\n", ". ", " ", ""]
            )
            
            chunks = text_splitter.split_text(content)
            
            # Create Document objects
            documents = []
            for i, chunk in enumerate(chunks):
                doc_metadata = {
                    "source": filename,
                    "file_type": "docx",
                    "chunk_index": i,
                    "upload_date": datetime.datetime.now().isoformat()
                }
                
                documents.append(Document(
                    page_content=chunk,
                    metadata=doc_metadata
                ))
            
            return documents
            
        except Exception as e:
            raise ValueError(f"Failed to process DOCX file: {str(e)}")
    
    def _process_document_by_type(self, file_path: str, file_type: str, filename: str) -> List[Document]:
        """Process document based on file type"""
        
        if file_type == 'pdf':
            return self._process_pdf(file_path, filename)
        elif file_type == 'docx':
            return self._process_docx(file_path, filename)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def process_all_documents(self) -> Dict:
        """Process all documents and create FAISS index"""
        print(f"\n🚀 DEBUG: === STARTING FAISS DOCUMENT PROCESSING ===")
        print(f"📊 DEBUG: Total documents to process: {len(self.uploaded_docs)}")
        
        # Clear existing vector store
        self._clear_vectorstore()
        
        processed_count = 0
        failed_count = 0
        total_chunks = 0
        failed_files = []
        
        # Collect all documents first
        all_documents = []
        
        for i, doc_info in enumerate(self.uploaded_docs):
            print(f"\n--- DEBUG: Processing document {i+1}/{len(self.uploaded_docs)} ---")
            print(f"📄 DEBUG: Filename: {doc_info['filename']}")
            
            try:
                file_path = doc_info["file_path"]
                
                if not Path(file_path).exists():
                    print(f"❌ DEBUG: File not found: {file_path}")
                    doc_info["status"] = "error"
                    failed_count += 1
                    failed_files.append(f"{doc_info['filename']} (file not found)")
                    continue
                
                # Process document
                documents = self._process_document_by_type(
                    file_path, 
                    doc_info["file_type"], 
                    doc_info["filename"]
                )
                print(f"📊 DEBUG: Generated {len(documents)} chunks")
                
                if len(documents) == 0:
                    print(f"⚠️ DEBUG: No chunks generated")
                    doc_info["status"] = "error"
                    failed_count += 1
                    failed_files.append(f"{doc_info['filename']} (no content extracted)")
                    continue
                
                # Add to collection
                all_documents.extend(documents)
                
                # Update tracking
                doc_info["chunk_count"] = len(documents)
                doc_info["status"] = "processed"
                doc_info["last_processed"] = datetime.datetime.now().isoformat()
                
                processed_count += 1
                total_chunks += len(documents)
                
            except Exception as e:
                print(f"❌ DEBUG: Failed to process {doc_info['filename']}: {e}")
                doc_info["status"] = "error"
                failed_count += 1
                failed_files.append(f"{doc_info['filename']} ({str(e)})")
        
        # Create FAISS index from all documents at once
        if all_documents:
            print(f"\n🗄️ DEBUG: Creating FAISS index from {len(all_documents)} chunks")
            try:
                # Create FAISS vector store
                vectorstore = FAISS.from_documents(
                    all_documents,
                    OpenAIEmbeddings()
                )
                
                # Save FAISS index
                vectorstore.save_local(str(self.vectorstore_path))
                print(f"✅ DEBUG: FAISS index saved successfully")
                
                # Verify files were created
                if self.faiss_index_path.exists() and self.faiss_metadata_path.exists():
                    print(f"✅ DEBUG: FAISS files verified on disk")
                else:
                    print(f"⚠️ DEBUG: FAISS files not found after save")
                
            except Exception as e:
                print(f"❌ DEBUG: Failed to create FAISS index: {e}")
                raise
        else:
            print(f"⚠️ DEBUG: No documents to index")
        
        # Save tracking
        self._save_document_tracking()
        
        result = {
            "processed_documents": processed_count,
            "failed_documents": failed_count,
            "total_chunks": total_chunks,
            "status": "completed"
        }
        
        if failed_files:
            result["failed_files"] = failed_files
        
        print(f"\n🏁 DEBUG: === FAISS PROCESSING COMPLETED ===")
        print(f"📊 DEBUG: Processed: {processed_count}")
        print(f"📊 DEBUG: Failed: {failed_count}")
        print(f"📊 DEBUG: Total chunks: {total_chunks}")
        
        return result
    
    def get_vectorstore(self):
        """Get FAISS vectorstore for querying"""
        try:
            if self.faiss_index_path.exists() and self.faiss_metadata_path.exists():
                return FAISS.load_local(
                    str(self.vectorstore_path),
                    OpenAIEmbeddings(),
                    allow_dangerous_deserialization=True
                )
            else:
                print(f"⚠️ DEBUG: FAISS index files not found")
                return None
        except Exception as e:
            print(f"❌ DEBUG: Failed to load FAISS index: {e}")
            return None
    
    def get_document_stats(self) -> Dict:
        """Get statistics about uploaded documents"""
        try:
            total_docs = len(self.uploaded_docs)
            total_chunks = sum(doc.get("chunk_count", 0) for doc in self.uploaded_docs)
            total_size = sum(doc.get("file_size", 0) for doc in self.uploaded_docs)
            
            # Count by status
            status_counts = {}
            file_type_counts = {}
            
            for doc in self.uploaded_docs:
                status = doc.get("status", "unknown")
                file_type = doc.get("file_type", "unknown")
                
                status_counts[status] = status_counts.get(status, 0) + 1
                file_type_counts[file_type] = file_type_counts.get(file_type, 0) + 1
            
            # Check if FAISS index exists
            index_exists = self.faiss_index_path.exists() and self.faiss_metadata_path.exists()
            
            return {
                "total_documents": total_docs,
                "total_chunks": total_chunks,
                "total_size_bytes": total_size,
                "status_breakdown": status_counts,
                "file_type_breakdown": file_type_counts,
                "index_exists": index_exists,
                "vectorstore_type": "FAISS",
                "documents": self.uploaded_docs
            }
        except Exception as e:
            logging.error(f"Failed to get document stats: {str(e)}")
            return {
                "total_documents": 0,
                "total_chunks": 0,
                "total_size_bytes": 0,
                "status_breakdown": {},
                "file_type_breakdown": {},
                "index_exists": False,
                "vectorstore_type": "FAISS",
                "documents": []
            }